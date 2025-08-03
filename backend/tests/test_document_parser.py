"""
文档解析模块测试
"""

import pytest
import tempfile
import sys
from pathlib import Path
from docx import Document
from openpyxl import Workbook

# 添加项目根目录到Python路径
sys.path.append(str(Path(__file__).parent.parent.parent))

from backend.app.services.document_parser import DocumentParser, ParsedDocument


@pytest.fixture
def temp_docs_dir():
    """创建临时文档目录和测试文件"""
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)

        # 创建测试Word文档
        doc = Document()
        doc.add_heading('测试文档标题', 0)
        doc.add_paragraph('这是第一段内容。')
        doc.add_heading('二级标题', level=2)
        doc.add_paragraph('这是第二段内容。')

        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '姓名'
        table.cell(0, 1).text = '年龄'
        table.cell(1, 0).text = '张三'
        table.cell(1, 1).text = '25'

        docx_path = temp_path / "test.docx"
        doc.save(docx_path)

        # 创建测试Excel文档
        wb = Workbook()
        ws = wb.active
        ws.title = "测试工作表"
        ws['A1'] = '姓名'
        ws['B1'] = '年龄'
        ws['A2'] = '李四'
        ws['B2'] = 30

        xlsx_path = temp_path / "test.xlsx"
        wb.save(xlsx_path)
        wb.close()

        # 创建测试文本文档
        txt_path = temp_path / "test.txt"
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("测试文本标题\n\n这是测试内容。\n\n重要说明:\n这是重要信息。")

        yield temp_path


def test_document_parser_init():
    """测试文档解析器初始化"""
    parser = DocumentParser()

    assert '.docx' in parser.supported_formats
    assert '.xlsx' in parser.supported_formats
    assert '.txt' in parser.supported_formats


def test_is_supported():
    """测试文件格式支持检查"""
    parser = DocumentParser()

    assert parser.is_supported("test.docx")
    assert parser.is_supported("test.xlsx")
    assert parser.is_supported("test.txt")
    assert not parser.is_supported("test.jpg")
    assert not parser.is_supported("test.pdf")


def test_parse_nonexistent_file():
    """测试解析不存在的文件"""
    parser = DocumentParser()

    result = parser.parse_document("/nonexistent/file.docx")

    assert not result.success
    assert "文件不存在" in result.error_message


def test_parse_unsupported_format():
    """测试解析不支持的格式"""
    parser = DocumentParser()

    # 创建临时文件
    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as f:
        temp_path = Path(f.name)

    try:
        result = parser.parse_document(temp_path)

        assert not result.success
        assert "不支持的文件格式" in result.error_message
    finally:
        temp_path.unlink()


def test_parse_docx(temp_docs_dir):
    """测试Word文档解析"""
    parser = DocumentParser()
    docx_path = temp_docs_dir / "test.docx"

    result = parser.parse_document(docx_path)

    assert result.success
    assert result.file_type == ".docx"
    assert "测试文档标题" in result.title
    assert "测试文档标题" in result.content
    assert "二级标题" in result.content
    assert "第一段内容" in result.content
    assert "第二段内容" in result.content

    # 检查表格转换
    assert "姓名" in result.content
    assert "张三" in result.content

    # 检查元数据
    assert 'paragraphs_count' in result.metadata
    assert 'tables_count' in result.metadata
    assert result.metadata['tables_count'] == 1


def test_parse_xlsx(temp_docs_dir):
    """测试Excel文档解析"""
    parser = DocumentParser()
    xlsx_path = temp_docs_dir / "test.xlsx"

    result = parser.parse_document(xlsx_path)

    assert result.success
    assert result.file_type == ".xlsx"
    assert "test" in result.title

    # 检查工作表标题
    assert "测试工作表" in result.content

    # 检查表格内容
    assert "姓名" in result.content
    assert "李四" in result.content
    assert "30" in result.content

    # 检查元数据
    assert 'sheets_count' in result.metadata
    assert 'sheet_names' in result.metadata
    assert result.metadata['sheets_count'] == 1


def test_parse_txt(temp_docs_dir):
    """测试文本文档解析"""
    parser = DocumentParser()
    txt_path = temp_docs_dir / "test.txt"

    result = parser.parse_document(txt_path)

    assert result.success
    assert result.file_type == ".txt"
    assert "测试文本标题" in result.title
    assert "测试内容" in result.content
    assert "重要说明" in result.content

    # 检查元数据
    assert 'lines_count' in result.metadata
    assert 'chars_count' in result.metadata
    assert result.metadata['lines_count'] > 0


def test_batch_parse(temp_docs_dir):
    """测试批量解析"""
    parser = DocumentParser()

    file_paths = [
        temp_docs_dir / "test.docx",
        temp_docs_dir / "test.xlsx",
        temp_docs_dir / "test.txt"
    ]

    results = parser.batch_parse(file_paths)

    assert len(results) == 3
    assert all(result.success for result in results)

    # 检查不同类型的文档都被正确解析
    file_types = [result.file_type for result in results]
    assert ".docx" in file_types
    assert ".xlsx" in file_types
    assert ".txt" in file_types


def test_parsed_document_structure():
    """测试ParsedDocument数据结构"""
    doc = ParsedDocument(
        file_path="/test/path.docx",
        file_name="path.docx",
        file_type=".docx",
        title="Test Title",
        content="Test Content",
        metadata={"key": "value"},
        parse_time=None
    )

    assert doc.file_path == "/test/path.docx"
    assert doc.file_name == "path.docx"
    assert doc.file_type == ".docx"
    assert doc.title == "Test Title"
    assert doc.content == "Test Content"
    assert doc.metadata == {"key": "value"}
    assert doc.success is True
    assert doc.error_message == ""


def test_markdown_formatting():
    """测试Markdown格式化"""
    parser = DocumentParser()

    # 测试文本格式化
    text = "TITLE\n\nThis is content.\n\nSection:\nMore content."
    formatted = parser._format_text_as_markdown(text)

    assert "## TITLE" in formatted
    assert "### Section:" in formatted
    assert "This is content." in formatted


def test_error_handling():
    """测试错误处理"""
    parser = DocumentParser()

    # 创建一个损坏的文件
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(b"invalid content")
        temp_path = Path(f.name)

    try:
        result = parser.parse_document(temp_path)

        # 应该捕获错误并返回失败结果
        assert not result.success
        assert result.error_message != ""
    finally:
        temp_path.unlink()
