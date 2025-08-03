"""
文档解析服务模块
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from loguru import logger


@dataclass
class ParsedDocument:
    """解析后的文档数据类"""
    file_path: str
    file_name: str
    file_type: str
    title: str
    content: str
    metadata: Dict
    parse_time: datetime
    success: bool = True
    error_message: str = ""


class DocumentParser:
    """文档解析器"""

    def __init__(self):
        """初始化文档解析器"""
        self.supported_formats = {
            '.docx': self._parse_docx,
            '.xlsx': self._parse_xlsx,
            '.txt': self._parse_txt
        }
        logger.info("文档解析器初始化完成")
        logger.info(f"支持格式: {list(self.supported_formats.keys())}")

    def parse_document(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        解析文档

        Args:
            file_path: 文档文件路径

        Returns:
            ParsedDocument: 解析结果
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return ParsedDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type="",
                title="",
                content="",
                metadata={},
                parse_time=datetime.now(),
                success=False,
                error_message="文件不存在"
            )

        file_type = file_path.suffix.lower()

        if file_type not in self.supported_formats:
            return ParsedDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_type,
                title="",
                content="",
                metadata={},
                parse_time=datetime.now(),
                success=False,
                error_message=f"不支持的文件格式: {file_type}"
            )

        logger.info(f"开始解析文档: {file_path}")

        try:
            parser_func = self.supported_formats[file_type]
            result = parser_func(file_path)
            logger.info(f"文档解析成功: {file_path}")
            return result

        except Exception as e:
            logger.error(f"文档解析失败 {file_path}: {e}")
            return ParsedDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_type,
                title="",
                content="",
                metadata={},
                parse_time=datetime.now(),
                success=False,
                error_message=str(e)
            )

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """解析Word文档"""
        doc = Document(file_path)

        # 提取标题
        title = ""
        if doc.paragraphs:
            title = doc.paragraphs[0].text.strip()

        # 提取内容
        content_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # 根据样式添加Markdown格式
            if paragraph.style.name.startswith('Heading'):
                level = self._extract_heading_level(paragraph.style.name)
                content_parts.append(f"{'#' * level} {text}")
            else:
                content_parts.append(text)

        # 处理表格
        for table in doc.tables:
            table_md = self._convert_table_to_markdown(table)
            content_parts.append(table_md)

        content = "\n\n".join(content_parts)

        # 提取元数据
        metadata = {
            'author': doc.core_properties.author or "",
            'created': doc.core_properties.created,
            'modified': doc.core_properties.modified,
            'subject': doc.core_properties.subject or "",
            'title': doc.core_properties.title or title,
            'paragraphs_count': len(doc.paragraphs),
            'tables_count': len(doc.tables)
        }

        return ParsedDocument(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=file_path.suffix.lower(),
            title=title,
            content=content,
            metadata=metadata,
            parse_time=datetime.now()
        )

    def _parse_xlsx(self, file_path: Path) -> ParsedDocument:
        """解析Excel文档"""
        workbook = load_workbook(file_path, data_only=True)

        # 使用文件名作为标题
        title = file_path.stem

        content_parts = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            # 添加工作表标题
            content_parts.append(f"## {sheet_name}")

            # 转换工作表为Markdown表格
            sheet_md = self._convert_worksheet_to_markdown(sheet)
            if sheet_md:
                content_parts.append(sheet_md)

        content = "\n\n".join(content_parts)

        # 提取元数据
        metadata = {
            'sheets_count': len(workbook.sheetnames),
            'sheet_names': workbook.sheetnames,
            'created': workbook.properties.created,
            'modified': workbook.properties.modified,
            'creator': workbook.properties.creator or "",
            'title': workbook.properties.title or title
        }

        workbook.close()

        return ParsedDocument(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=file_path.suffix.lower(),
            title=title,
            content=content,
            metadata=metadata,
            parse_time=datetime.now()
        )

    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        """解析文本文档"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # 使用第一行作为标题
        lines = content.split('\n')
        title = lines[0].strip() if lines else file_path.stem

        # 基本的Markdown格式化
        content = self._format_text_as_markdown(content)

        # 获取文件统计信息
        stat = file_path.stat()
        metadata = {
            'size': stat.st_size,
            'created': datetime.fromtimestamp(stat.st_ctime),
            'modified': datetime.fromtimestamp(stat.st_mtime),
            'lines_count': len(lines),
            'chars_count': len(content)
        }

        return ParsedDocument(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=file_path.suffix.lower(),
            title=title,
            content=content,
            metadata=metadata,
            parse_time=datetime.now()
        )

    def _extract_heading_level(self, style_name: str) -> int:
        """从样式名称提取标题级别"""
        match = re.search(r'(\d+)', style_name)
        return int(match.group(1)) if match else 1

    def _convert_table_to_markdown(self, table) -> str:
        """将Word表格转换为Markdown格式"""
        if not table.rows:
            return ""

        rows = []

        # 处理表头
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        rows.append("| " + " | ".join(headers) + " |")
        rows.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # 处理数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        return "\n".join(rows)

    def _convert_worksheet_to_markdown(self, sheet: Worksheet) -> str:
        """将Excel工作表转换为Markdown格式"""
        if sheet.max_row == 0 or sheet.max_column == 0:
            return ""

        rows = []

        # 获取数据范围
        data_rows = list(sheet.iter_rows(values_only=True))

        if not data_rows:
            return ""

        # 处理表头
        headers = [str(cell) if cell is not None else "" for cell in data_rows[0]]
        rows.append("| " + " | ".join(headers) + " |")
        rows.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # 处理数据行
        for row_data in data_rows[1:]:
            cells = [str(cell) if cell is not None else "" for cell in row_data]
            # 确保列数一致
            while len(cells) < len(headers):
                cells.append("")
            rows.append("| " + " | ".join(cells[:len(headers)]) + " |")

        return "\n".join(rows)

    def _format_text_as_markdown(self, text: str) -> str:
        """将纯文本格式化为基本的Markdown"""
        lines = text.split('\n')
        formatted_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                formatted_lines.append("")
                continue

            # 简单的格式检测
            if line.isupper() and len(line) < 100:
                # 全大写短行可能是标题
                formatted_lines.append(f"## {line}")
            elif line.endswith(':') and len(line) < 100:
                # 以冒号结尾的可能是小标题
                formatted_lines.append(f"### {line}")
            else:
                formatted_lines.append(line)

        return "\n".join(formatted_lines)

    def batch_parse(self, file_paths: List[Union[str, Path]]) -> List[ParsedDocument]:
        """批量解析文档"""
        results = []

        logger.info(f"开始批量解析 {len(file_paths)} 个文档")

        for file_path in file_paths:
            result = self.parse_document(file_path)
            results.append(result)

        success_count = sum(1 for r in results if r.success)
        logger.info(f"批量解析完成: 成功 {success_count}/{len(file_paths)}")

        return results

    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """检查文件格式是否支持"""
        file_path = Path(file_path)
        return file_path.suffix.lower() in self.supported_formats
